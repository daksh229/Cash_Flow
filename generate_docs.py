"""
Generate two client-ready DOCX documents:
  1. System_Architecture.docx
  2. Project_Plan.docx

Rules followed:
  - Aligned with CashFlow_SDD_v17 (Phase 1 scope, 9 modules, shared Feature Store,
    S7 single source of truth, Recommendation Engine with feedback).
  - No file names, folder paths, or code references.
  - No budgets, no approval tables, no sign-off blocks.
  - No symbols like star / check / warning / emoji.
  - Specific wording. Tables for structured content.
  - ASCII-free diagrams rendered as block rectangles with flow lines.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------------------------------------------------
# Styling helpers
# --------------------------------------------------------------------

def set_cell_bg(cell, colour_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), colour_hex)
    tc_pr.append(shd)


def add_border(paragraph, colour="BFBFBF"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'bottom', 'left', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:color'), colour)
        pBdr.append(b)
    pPr.append(pBdr)


def style_body(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return h


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    return p


def styled_table(doc, header_row, data_rows, col_widths=None,
                 header_bg="0D47A1", stripe_bg="EAF2FB"):
    tbl = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header_row):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_bg(hdr[i], header_bg)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(data_rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_bg(cell, stripe_bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl


def page_break(doc):
    doc.add_page_break()


def callout(doc, label, body, colour="EAF2FB"):
    """A coloured box with a bold label and body text."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, colour)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph()
    return tbl


# ====================================================================
# DOCUMENT 1 — SYSTEM ARCHITECTURE
# ====================================================================

def build_system_architecture():
    doc = Document()
    style_body(doc)

    # Cover
    title = doc.add_heading('Cash Flow Forecasting Platform', level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    sub = doc.add_paragraph()
    run = sub.add_run('System Design and Architecture')
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    run = meta.add_run(
        'Aligned with the Cash Flow Forecasting Solution Design Document '
        '(Phase 1, Algorithm Layer).'
    )
    run.italic = True
    run.font.size = Pt(11)

    page_break(doc)

    # 1. Executive Summary
    heading(doc, '1. Executive Summary', level=1)
    para(doc,
         'The Cash Flow Forecasting Platform is a shared-intelligence financial system. '
         'It unifies receivables, payables, projects, sales pipeline, contingent inflows '
         'and expenses into a single daily cash position and produces ranked treasury '
         'recommendations that learn from outcomes.')
    para(doc,
         'The design is built around four architectural commitments:')
    bullet(doc, 'A shared feature store is the single source of behavioural truth for every model.')
    bullet(doc, 'Rule-based forecasting is the deterministic baseline; machine learning is additive, never foundational.')
    bullet(doc, 'The system is event-reactive: business changes update the forecast continuously, not only at nightly batch.')
    bullet(doc, 'Every cash event, recommendation, and model prediction is traceable through audit and lineage records.')

    # 2. Scope and Phase Boundary
    heading(doc, '2. Scope and Phase Boundary', level=1)
    para(doc,
         'Phase 1 covers the algorithm layer end-to-end: six forecasting modules, '
         'one aggregation module, and a scoring-based recommendation engine. '
         'Probabilistic forecasting, Monte Carlo simulation, reinforcement-learning feedback, '
         'and real-time explainability are planned for later phases.')

    heading(doc, '2.1 In Scope (Phase 1)', level=2)
    bullet(doc, 'S1 AR Collections Prediction with dynamic, event-driven re-scoring.')
    bullet(doc, 'S2 Vendor Payment Prediction including treasury liquidity gate and vendor prioritisation.')
    bullet(doc, 'Credit Risk Classification into LOW, MEDIUM, and HIGH bands.')
    bullet(doc, 'S3 WIP and milestone billing forecast (deterministic rule-based).')
    bullet(doc, 'S4 Sales pipeline forecast (cohort-based deterministic).')
    bullet(doc, 'S5 Contingent inflows forecast (scheduling).')
    bullet(doc, 'S6 Expense forecast across salary, tax, recurring, and non-PO categories.')
    bullet(doc, 'S7 Cash event normalisation, trust scoring, deduplication, and aggregation.')
    bullet(doc, 'Recommendation Engine across collections, vendor deferral, and expense deferral levers.')
    bullet(doc, 'Shared feature store with versioning.')
    bullet(doc, 'Reconciliation of forecast against realised outcomes and composite cash-accuracy reporting.')
    bullet(doc, 'Multi-entity deployment with tenant-scoped persistence.')

    heading(doc, '2.2 Out of Scope (Phase 2 and Beyond)', level=2)
    bullet(doc, 'Probability curves for payment timing (survival / hazard models).')
    bullet(doc, 'Monte Carlo simulation producing confidence intervals on the forecast.')
    bullet(doc, 'Reinforcement-learning feedback loop with automatic weight promotion.')
    bullet(doc, 'Multi-objective optimisation across cash, risk, and operational constraints.')
    bullet(doc, 'Real-time treasury cockpit with a full explainability layer.')
    bullet(doc, 'Direct ERP or CRM connectors. Raw source data is consumed only through the Data Hub.')

    # 3. Architectural Principles
    heading(doc, '3. Architectural Principles', level=1)
    styled_table(doc,
        ['Principle', 'Intent'],
        [
            ['Deterministic first, ML additive',
             'Rule-based modules produce the Phase 1 baseline. Machine learning refines, but never substitutes, the deterministic answer.'],
            ['Shared feature store as single source of truth',
             'Every model reads behavioural features from the same versioned store. No module computes features from raw data at runtime.'],
            ['Event-reactive, not cron-first',
             'Business events trigger immediate re-scoring. Batch runs consolidate; they do not drive the system.'],
            ['Multi-entity by construction',
             'Every persisted row is scoped to an entity identifier so one deployment serves multiple business entities safely.'],
            ['Non-destructive audit',
             'Suppressed events and replaced records are retained with reason codes. Nothing is deleted silently.'],
            ['Configurable policy, not hardcoded logic',
             'Vendor priority, scoring weights, stage probabilities, thresholds, and escalation rules live in configuration tables controlled by treasury.'],
            ['Explainable recommendations',
             'Every recommendation states what to do, why, the entity involved, and the expected cash impact.'],
        ],
        col_widths=[5, 11])

    # 4. Logical Architecture
    heading(doc, '4. Logical Architecture', level=1)
    para(doc,
         'The platform is organised into five logical layers. Each layer has a single '
         'responsibility and depends only on layers below it. No layer is allowed to '
         'skip or bypass another.')

    styled_table(doc,
        ['Layer', 'Responsibility', 'Key Services'],
        [
            ['Presentation',
             'Serves treasury users and external consumers.',
             'Application programming interface, dashboard, captured-expense form, health and metrics endpoints.'],
            ['Domain',
             'Forecasting, aggregation, recommendation, reconciliation.',
             'S1 to S7, Recommendation Engine, reconciliation engine, composite cash-accuracy computation.'],
            ['Orchestration',
             'Decides what runs, when, in what order.',
             'Dependency-aware task runner, event bus with persistence, volume-triggered retraining.'],
            ['Integration',
             'Talks to the outside world.',
             'Inbound event adapter with idempotency and dead letter queue, outbound publisher.'],
            ['Infrastructure',
             'Persistence, security, observability, cross-cutting primitives.',
             'Database layer, feature store with version policy, authentication and role-based access control, audit and lineage store, metrics and logging.'],
        ],
        col_widths=[3.2, 5.5, 7.3])

    heading(doc, '4.1 Layered View (Schematic)', level=2)
    para(doc,
         'The diagram below expresses the layered dependency model. Arrows indicate '
         'read or call direction. No upward calls are permitted.')

    diag_table = doc.add_table(rows=5, cols=1)
    diag_table.style = 'Light Grid Accent 1'
    layer_colours = ["D9E8F7", "CFE4D9", "FFF2CC", "FCE5CD", "E6E6E6"]
    layer_labels = [
        ("Presentation", "Application programming interface  |  Dashboard  |  Captured-expense form  |  Health and metrics"),
        ("Domain", "S1  |  S2  |  Credit Risk  |  S3  |  S4  |  S5  |  S6  |  S7  |  Recommendation Engine  |  Reconciliation"),
        ("Orchestration", "Task runner  |  Event bus with persistence  |  Volume-triggered retraining"),
        ("Integration", "Inbound event adapter  |  Idempotency and dead letter queue  |  Outbound publisher"),
        ("Infrastructure", "Database  |  Feature store with version policy  |  Audit and lineage  |  Security  |  Monitoring"),
    ]
    for i, (name, desc) in enumerate(layer_labels):
        cell = diag_table.rows[i].cells[0]
        set_cell_bg(cell, layer_colours[i])
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{name}\n")
        r.bold = True
        r.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.size = Pt(9)

    # 5. Central Financial Intelligence Layer
    heading(doc, '5. Central Financial Intelligence Layer', level=1)
    para(doc,
         'All forecasting modules consume behavioural features from a shared, persisted store. '
         'Customer and vendor behaviour — payment scores, delays, volatility, discount patterns, '
         'advance-payment history — is pre-computed and reused across every module. '
         'No module recomputes features from raw data at prediction time.')

    heading(doc, '5.1 Feature Store Contents', level=2)
    styled_table(doc,
        ['Feature Set', 'Entity', 'Primary Signals'],
        [
            ['Customer features', 'Customer',
             'Average and median payment delay, late-payment ratio, payment volatility, dispute ratio, recent reminder velocity, open dispute count, days sales outstanding, seasonality, advance-payment ratio.'],
            ['Customer payment scores', 'Customer',
             'Reliability index, expected delay for next invoice, risk segment.'],
            ['Vendor features', 'Vendor',
             'Average payment cycle, payment volatility, discount capture ratio, chase frequency, purchase-order to bill lag, advance-payment ratio.'],
            ['Invoice features', 'Transaction',
             'Invoice age, days past due, amount bucket, amount percentile within customer, partial-payment flag.'],
            ['Bill features', 'Transaction',
             'Bill age, days past due, amount bucket, approval status, early-payment eligibility, penalty accrual flag.'],
            ['Collections features', 'Transaction',
             'Reminder count, call count, promise-to-pay flag, days since last contact, promise-kept ratio.'],
        ],
        col_widths=[4.5, 2.8, 8.7])

    heading(doc, '5.2 Feature Lifecycle', level=2)
    styled_table(doc,
        ['Stage', 'Trigger', 'Outcome'],
        [
            ['Compute', 'Nightly batch plus lifecycle events',
             'Features written to the persisted store, tagged with an entity identifier, a feature-set name, and a deterministic version string.'],
            ['Register', 'Every write',
             'A new version is recorded in draft state with row count and configuration hash.'],
            ['Promote', 'Explicit operator action',
             'One version per feature set becomes active; previous active versions are retired.'],
            ['Freeze', 'When a downstream model or forecast depends on a version',
             'Version is held read-only for reproducibility of past predictions.'],
            ['Retire', 'Superseded by a newer active version',
             'Version remains readable but is not selected by default readers.'],
        ],
        col_widths=[2.5, 5.5, 8])

    # 6. Forecasting Modules
    heading(doc, '6. Forecasting Modules', level=1)
    para(doc,
         'Nine modules span both cash inflows and outflows. Each module is '
         'independently runnable, reads from the shared feature store, and writes '
         'to a common cash-event output table.')

    heading(doc, '6.1 Module Catalogue', level=2)
    styled_table(doc,
        ['Code', 'Name', 'Method', 'Primary Output'],
        [
            ['S1', 'AR Collections Prediction', 'Gradient boosting with a random-forest baseline',
             'Predicted payment date for every open invoice.'],
            ['S2', 'Vendor Payment Prediction', 'Five-layer decision pipeline with a machine-learning adjustment and liquidity gate',
             'Final scheduled payment date for every approved bill.'],
            ['Credit Risk', 'Customer Risk Classification', 'Multiclass classifier',
             'Risk band of LOW, MEDIUM, or HIGH for each customer.'],
            ['S3', 'WIP and Project Milestone Billing', 'Deterministic rule pipeline',
             'Expected invoice and cash dates for near-complete milestones.'],
            ['S4', 'Sales Pipeline Forecast', 'Cohort matching with probability-weighted deal value',
             'Expected cash from open opportunities.'],
            ['S5', 'Contingent Inflows Forecast', 'Scheduling by approval status',
             'Expected cash from loans, grants, refunds, insurance, and other non-trading sources.'],
            ['S6', 'Expense Forecast', 'Category-based scheduling including non-PO capture',
             'Expected outflow schedule across salary, tax, recurring, PO-based, and ad-hoc expenses.'],
            ['S7', 'Cash Event Aggregation', 'Normalise, trust-score, deduplicate, audit',
             'Unified daily, weekly, and monthly cash position.'],
            ['Recommendation Engine', 'Treasury Action Ranking',
             'Multi-dimension scoring and ranking with feedback capture',
             'Ranked, explainable recommendations across collections, vendor deferral, and expense deferral.'],
        ],
        col_widths=[1.5, 4.5, 5.5, 5.5])

    heading(doc, '6.2 Module Dependencies', level=2)
    para(doc,
         'The dependency graph is declarative. The task runner builds execution order '
         'from this graph and parallelises independent branches.')

    styled_table(doc,
        ['Task', 'Depends On'],
        [
            ['Feature table generation', 'None'],
            ['S1 AR Prediction', 'Feature table generation'],
            ['S2 AP Prediction', 'Feature table generation'],
            ['Credit Risk', 'Feature table generation'],
            ['S3 WIP Forecast', 'Feature table generation'],
            ['S4 Pipeline Forecast', 'Feature table generation'],
            ['S5 Contingent Inflows', 'None'],
            ['S6 Expense Forecast', 'None'],
            ['S7 Cash Aggregation', 'S1, S2, S3, S4, S5, S6'],
            ['Recommendation Engine', 'S7, Credit Risk'],
        ],
        col_widths=[6.5, 10])

    # 7. S1 in detail
    heading(doc, '7. S1 Detailed Design', level=1)
    para(doc,
         'S1 predicts the expected payment date for every open receivable invoice '
         'using machine learning. The prediction refreshes on every meaningful lifecycle '
         'event rather than only during nightly batch.')

    heading(doc, '7.1 Prediction Target', level=2)
    bullet(doc, 'Primary target: days to pay, measured from invoice date to observed payment date.')
    bullet(doc, 'Derived output: predicted payment date, computed as invoice date plus predicted days to pay.')
    bullet(doc, 'Confidence tier: HIGH, MEDIUM, or LOW based on model agreement and customer history depth.')

    heading(doc, '7.2 Model Strategy', level=2)
    styled_table(doc,
        ['Model Role', 'Description'],
        [
            ['Primary', 'Gradient boosting regression. Handles mixed numerical and categorical features. Produces the headline prediction.'],
            ['Baseline', 'Random forest regression. Runs in parallel. Large divergence from the primary output flags uncertainty for human review.'],
            ['Prior', 'Three-level hierarchical prior combining customer-level, segment-level, and global-level statistics using empirical-Bayes shrinkage. Used when a customer has less history than the configured minimum.'],
        ],
        col_widths=[3, 13])

    heading(doc, '7.3 Lifecycle Triggers', level=2)
    para(doc,
         'The following lifecycle events cause an immediate re-score of all affected invoices. '
         'Triggers also update rolling features at the customer level so predictions for any '
         'other open invoice for the same customer reflect the latest behaviour.')
    styled_table(doc,
        ['Event', 'Effect'],
        [
            ['Invoice created', 'First prediction made using the customer score plus invoice attributes.'],
            ['Invoice viewed on portal', 'Strong positive signal. Predicted delay reduced.'],
            ['Reminder sent', 'Reminder count incremented. Re-score.'],
            ['Promise to pay received', 'Promise date used as an anchor. Re-score with lower risk weight.'],
            ['Dispute raised', 'Risk increases. Predicted delay extended. Re-score.'],
            ['Dispute resolved', 'Risk partially restored. Re-score.'],
            ['Partial payment made', 'Outstanding balance reduced. Promise-kept ratio updated.'],
            ['Days past due threshold', 'Scheduled re-score at seven, fifteen, thirty, and sixty days overdue.'],
            ['Payment received in full', 'Invoice closed. Actual days to pay recorded for retraining.'],
        ],
        col_widths=[4.5, 11.5])

    heading(doc, '7.4 Serving-Time Decision', level=2)
    para(doc,
         'At serving time, a model-selection step chooses between the primary, baseline, '
         'and prior outputs based on three inputs:')
    bullet(doc, 'Per-customer history depth compared against the configured thin-data threshold.')
    bullet(doc, 'Recent metric history of the primary model compared against a degradation threshold.')
    bullet(doc, 'The active variant recorded in the model registry for the serving entity.')
    para(doc,
         'If the primary output is degraded beyond the threshold, the baseline is used. '
         'If history is below the minimum required for the primary, the prior is used. '
         'If the model registry records a different active variant, that variant is honoured.')

    # 8. S2 Vendor Payments
    heading(doc, '8. S2 Detailed Design', level=1)
    para(doc,
         'S2 predicts the exact date each vendor obligation will result in a cash outflow. '
         'It combines policy, machine learning, and a liquidity check in a five-layer '
         'sequential pipeline.')

    styled_table(doc,
        ['Layer', 'Description'],
        [
            ['Earliest payable date', 'Checks that the bill is approved, payment terms are satisfied, and the due date has been reached.'],
            ['Payment run alignment', 'Snaps the payment to the next eligible payment run after the earliest payable date.'],
            ['Vendor prioritisation', 'Applies the treasury-owned priority table. Strategic, high-penalty, and early-discount vendors move to the front of the queue.'],
            ['Machine-learning adjustment', 'Predicts an adjustment in days against the rule-based candidate date, learned from historical patterns.'],
            ['Treasury liquidity gate', 'Checks the forecast cash position on the planned payment date. Defers lower-priority payments when the balance falls below the configured floor.'],
        ],
        col_widths=[4.5, 11.5])

    heading(doc, '8.1 Vendor Prioritisation Dimensions', level=2)
    styled_table(doc,
        ['Dimension', 'Trade-off'],
        [
            ['Priority tier', 'Tier one vendors are paid first regardless of liquidity. This can reduce cash available for lower tiers.'],
            ['Early-payment discount', 'Compares discount value against the cost of reduced liquidity. Pay early only when the net value is positive.'],
            ['Late-payment penalty', 'Penalty exceeding the benefit of deferring keeps the bill on schedule. High-penalty vendors self-prioritise.'],
            ['Relationship sensitivity', 'High-relationship vendors are paid ahead of terms to protect the supply chain. This is a policy choice, not a pure financial calculation.'],
            ['Payment run cadence', 'A bill due on a day without a payment run waits to the next run. The model accounts for this lag.'],
            ['Liquidity gate override', 'Below the cash threshold, lower-priority bills defer. Tier one deferral escalates to treasury for manual decision.'],
        ],
        col_widths=[4, 12])

    # 9. S7 Aggregation
    heading(doc, '9. S7 Cash Event Aggregation', level=1)
    para(doc,
         'S7 is the single source of truth for the downstream cash position. '
         'It executes a four-stage pipeline:')
    numbered(doc,
             'Normalisation: convert every upstream output into a canonical event '
             'with entity identifier, event date, amount with sign, direction, '
             'confidence, and currency.')
    numbered(doc,
             'Trust scoring: assign a source-specific trust baseline based on the '
             'module type (deterministic sources receive higher baselines than '
             'machine-learning sources), adjusted by recent accuracy metrics.')
    numbered(doc,
             'Deduplication: bucket events by entity, rounded amount, and a date '
             'window. Within each bucket the highest-trust event wins and the '
             'suppressed events are recorded with reason codes.')
    numbered(doc,
             'Audit: write a per-run summary with inputs, kept and dropped counts, '
             'and a lineage edge linking the output to every contributing source.')

    heading(doc, '9.1 Source Trust Baselines', level=2)
    styled_table(doc,
        ['Source Module', 'Baseline', 'Rationale'],
        [
            ['S5 Contingent Inflows', 'High',
             'Deterministic source driven by confirmed commitments and approval status.'],
            ['S6 Expense Forecast', 'High',
             'Deterministic scheduling; salary, tax, and fixed recurring entries are known with high certainty.'],
            ['S2 AP Prediction', 'High-Medium',
             'Rule pipeline with a small machine-learning adjustment and explicit liquidity check.'],
            ['S1 AR Prediction', 'Medium',
             'Machine-learning prediction; trust rises with customer history depth and recent model accuracy.'],
            ['S3 WIP Forecast', 'Medium',
             'Deterministic from project milestones; lag estimate introduces some uncertainty.'],
            ['S4 Pipeline Forecast', 'Lower-Medium',
             'Probabilistic deal conversion; highest inherent variance among Phase 1 modules.'],
        ],
        col_widths=[4.5, 2.5, 9])

    # 10. Recommendation Engine
    heading(doc, '10. Recommendation Engine', level=1)
    para(doc,
         'The Recommendation Engine is the central decision layer. It reads the aggregated '
         'cash position and the customer intelligence store, generates candidate scenarios '
         'across three levers, scores and ranks them, and captures user feedback to learn '
         'over time.')

    heading(doc, '10.1 Scoring Dimensions', level=2)
    styled_table(doc,
        ['Dimension', 'Description'],
        [
            ['Cash improvement', 'Expected positive change to the cash position over the recommendation horizon.'],
            ['Risk reduction', 'Reduction in exposure to a specific customer, vendor, or category as a result of the action.'],
            ['Target alignment', 'Consistency with the treasury policy and target operating balances.'],
            ['Feasibility', 'Operational and policy-level feasibility, including priority-tier protection and minimum-balance constraints.'],
        ],
        col_widths=[4, 12])

    heading(doc, '10.2 Levers', level=2)
    styled_table(doc,
        ['Lever', 'Typical Actions'],
        [
            ['Collections acceleration', 'Issue reminders, request promise-to-pay, escalate disputes, offer settlement discounts where appropriate.'],
            ['Vendor deferral', 'Push eligible lower-priority bills to a later payment run, honour discounts on higher-priority bills, protect tier-one vendors.'],
            ['Expense deferral', 'Defer eligible non-critical operational expenses such as seasonal or one-time categories within policy limits.'],
        ],
        col_widths=[4, 12])

    heading(doc, '10.3 Feedback Loop', level=2)
    para(doc,
         'Every recommendation surfaced to the user can be accepted, rejected, or ignored. '
         'Accepted recommendations are later paired with the realised cash impact observed '
         'via reconciliation. A weight-tuning step reads the accumulated sample and proposes '
         'updated scoring weights. Proposals are advisory; promotion into production is an '
         'explicit operator action.')

    # 11. Reconciliation and KPI
    heading(doc, '11. Reconciliation and Cash Accuracy', level=1)
    para(doc,
         'Reconciliation joins forecasted cash events with realised outcomes ingested from '
         'the Data Hub. Variance is computed at invoice, bill, and aggregate levels.')

    heading(doc, '11.1 Metrics Produced', level=2)
    styled_table(doc,
        ['Metric', 'Definition'],
        [
            ['Match rate', 'Share of forecasted events with a corresponding realised outcome.'],
            ['Mean absolute error in days', 'Average number of days between forecast date and actual payment date among matched events.'],
            ['Bias in days', 'Signed mean error; positive indicates forecasts were late on average.'],
            ['Mean absolute percentage error on amount', 'Average absolute percentage difference between forecast and actual amounts among matched events.'],
        ],
        col_widths=[5, 11])

    heading(doc, '11.2 Composite Cash Accuracy', level=2)
    para(doc,
         'The composite cash-accuracy indicator combines amount-level and date-level accuracy '
         'with cash-accuracy weighted higher than date-accuracy. Weights and the days target '
         'are operator-configurable so the business can express its priorities explicitly.')
    styled_table(doc,
        ['Component', 'Formula (verbal)', 'Default Weight'],
        [
            ['Cash accuracy', 'One minus the amount error, clipped to the range zero to one, expressed as a percentage.', '70 percent'],
            ['Days accuracy', 'One minus the date error divided by the configured days target, floored at zero, expressed as a percentage.', '30 percent'],
            ['Composite indicator', 'Weighted sum of the two.', 'Computed'],
        ],
        col_widths=[3.5, 9.5, 3])

    # 12. Event-Driven Operation
    heading(doc, '12. Event-Driven Operation', level=1)
    para(doc,
         'The platform operates as an event-reactive system. Business events flow from '
         'the external Data Hub through an inbound adapter, onto an internal event bus, '
         'and into the appropriate partial re-run of the task graph.')

    heading(doc, '12.1 Events Handled', level=2)
    styled_table(doc,
        ['Event', 'Response'],
        [
            ['Invoice created, paid, updated', 'Partial re-run of the receivables prediction branch for the affected entity.'],
            ['Bill created, paid, updated', 'Partial re-run of the payables prediction branch for the affected entity.'],
            ['Customer or vendor profile updated', 'Feature table rebuild cascading into every consuming module.'],
            ['Forecast published', 'Outbound publication of the latest forecast summary to the Data Hub.'],
        ],
        col_widths=[5, 11])

    heading(doc, '12.2 Volume-Triggered Retraining', level=2)
    para(doc,
         'A counter keyed by entity and model records inbound event volume. When the counter '
         'reaches a configured threshold, a full re-run for that branch is triggered and the '
         'counter resets. This replaces fixed retraining cadences with volume-based retraining.')

    # 13. Integration
    heading(doc, '13. Data Hub Integration', level=1)
    heading(doc, '13.1 Inbound', level=2)
    bullet(doc, 'Signed webhook accepting canonical envelopes. Every request is verified against a shared signing key before dispatch.')
    bullet(doc, 'Idempotency enforced by a persisted envelope-identifier ledger. Duplicate envelopes are acknowledged without re-processing.')
    bullet(doc, 'Malformed or unmapped envelopes are written to a dead letter queue with reason and error for later inspection and replay.')
    bullet(doc, 'Bulk endpoint and a command-line replay for historical backfills.')
    heading(doc, '13.2 Outbound', level=2)
    bullet(doc, 'Signed publication of forecast summaries back to the Data Hub for consumption by other platform components.')
    bullet(doc, 'Local append-only fallback when the Data Hub endpoint is unavailable; ensures no publication is lost.')

    # 14. Multi-Entity
    heading(doc, '14. Multi-Entity Support', level=1)
    bullet(doc, 'Every persisted record carries an entity identifier.')
    bullet(doc, 'The active entity is resolved through a request-scoped context for online requests, and through an environment variable for batch jobs.')
    bullet(doc, 'All registry lookups, event emissions, audit rows, and forecast outputs are filtered by the active entity.')
    bullet(doc, 'Automated tests verify that queries issued under one entity cannot see data from another.')

    # 15. Security
    heading(doc, '15. Security, Access, and Secrets', level=1)
    bullet(doc, 'Stateless bearer tokens signed with a shared key, carrying subject, roles, issued-at, and expiry.')
    bullet(doc, 'Three coarse roles: Viewer, Analyst, and Administrator, in a strict inclusion order.')
    bullet(doc, 'Secrets are resolved from environment variables, a local environment file, and a mounted secret directory, in that order. Secrets are never stored in the configuration document.')
    bullet(doc, 'All write endpoints are role-gated. All inbound integration traffic is signature-verified.')

    # 16. Observability
    heading(doc, '16. Observability', level=1)
    heading(doc, '16.1 Metrics', level=2)
    styled_table(doc,
        ['Metric Family', 'What it Measures'],
        [
            ['Run counters', 'Successful and failed task runs by pipeline.'],
            ['Run durations', 'End-to-end time per pipeline run.'],
            ['Model accuracy gauges', 'Latest accuracy indicator per model and per entity.'],
            ['Event throughput', 'Emitted events per event name.'],
            ['Integration errors', 'Count of dead letter rows by reason.'],
        ],
        col_widths=[5, 11])
    heading(doc, '16.2 Health and Logging', level=2)
    bullet(doc, 'Liveness endpoint: returns an acknowledgement that the process is running.')
    bullet(doc, 'Readiness endpoint: verifies database reachability before accepting traffic.')
    bullet(doc, 'Structured logging with a machine-readable format in production and a human-readable format in development.')
    bullet(doc, 'Every log line from a task run carries the run identifier so lines can be grouped end to end.')

    # 17. Reliability
    heading(doc, '17. Reliability Primitives', level=1)
    bullet(doc, 'Domain exception hierarchy separating configuration errors, data-validation errors, upstream-data-missing errors, training errors, and transient external-service errors.')
    bullet(doc, 'Retry policy with exponential backoff and jitter for transient external-service errors.')
    bullet(doc, 'Circuit breaker protecting external dependencies. Closed, open, and half-open states with configurable failure threshold and reset interval.')

    # 18. Data Model
    heading(doc, '18. Data Model Summary', level=1)
    para(doc,
         'The persisted data model is a small set of tables, each scoped to an entity identifier. '
         'Everything else is derived.')

    styled_table(doc,
        ['Table', 'Purpose'],
        [
            ['Feature snapshots', 'Versioned feature rows consumed by every model.'],
            ['Feature versions', 'State machine for feature versions (draft, active, frozen, retired).'],
            ['Forecast outputs', 'Unified cash events produced by every module with reference identifiers for matching.'],
            ['Run audit', 'One row per task graph run with status, start and end timestamps, and error reason when applicable.'],
            ['Event log', 'Every emitted event with persistence before dispatch. Supports replay.'],
            ['Actual outcomes', 'Realised cash events ingested from the Data Hub. Input to reconciliation.'],
            ['Recommendation feedback', 'User accept, reject, or ignore action, and the realised cash impact when available.'],
            ['Model registry', 'Per-entity, per-model, per-variant serving state (active, shadow, retired).'],
            ['Captured non-PO expenses', 'Manually entered operational expenses that do not pass through a purchase order.'],
            ['Integration dead letter queue', 'Envelopes that could not be mapped or handled, preserved with reason for later replay.'],
            ['Integration idempotency ledger', 'Envelope identifiers seen within a retention window.'],
        ],
        col_widths=[5, 11])

    # 19. Deployment
    heading(doc, '19. Deployment Profile', level=1)
    bullet(doc, 'Containerised services running as a non-root process with a restricted filesystem.')
    bullet(doc, 'Application programming interface and dashboard served as separate containers.')
    bullet(doc, 'Relational database for production; file-based database for development.')
    bullet(doc, 'Secrets mounted from a secrets directory rather than environment variables in production.')
    bullet(doc, 'Kubernetes manifests providing liveness and readiness probes, resource requests and limits, and drop-all capability policy.')

    # 20. Design Principles Recap
    heading(doc, '20. Summary', level=1)
    para(doc,
         'The architecture is intentionally simple at the top and deep where it matters. '
         'A small set of layered services. A shared intelligence layer that every model '
         'consumes. A deterministic baseline that is refined, not replaced, by machine '
         'learning. A single source of truth in the aggregation step. A feedback loop that '
         'learns from outcomes. Multi-entity from day one. Every decision traceable through '
         'audit and lineage.')

    out = "System_Architecture.docx"
    doc.save(out)
    print("wrote", out)


# ====================================================================
# DOCUMENT 2 — PROJECT PLAN
# ====================================================================

def build_project_plan():
    doc = Document()
    style_body(doc)

    title = doc.add_heading('Cash Flow Forecasting Platform', level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    sub = doc.add_paragraph()
    run = sub.add_run('Project Plan — Real-Data Rollout')
    run.font.size = Pt(18)
    run.bold = True

    meta = doc.add_paragraph()
    run = meta.add_run('Sprint-by-sprint delivery plan for the transition from '
                       'synthetic-data baseline to multi-entity production rollout.')
    run.italic = True
    run.font.size = Pt(11)

    page_break(doc)

    # 1. Overview
    heading(doc, '1. Overview', level=1)
    para(doc,
         'The plan spans twenty calendar weeks across eight sprints. Sprint cadence is two '
         'weeks by default. Three sprints carry additional buffer where the risk profile '
         'is highest: the first sprint absorbing integration uncertainty, the sixth sprint '
         'dedicated to the recommendation engine and user acceptance, and the seventh '
         'sprint covering hardening and production cut-over.')

    styled_table(doc,
        ['Phase', 'Sprints', 'Weeks'],
        [
            ['Real-data integration', 'Sprint 1', '3'],
            ['Real-data preprocessing', 'Sprint 2', '2'],
            ['Machine-learning retraining on real data', 'Sprint 3', '2'],
            ['Rule-based and aggregation calibration', 'Sprint 4', '2'],
            ['Reconciliation and cash-accuracy indicator', 'Sprint 5', '2'],
            ['Recommendation engine calibration and user acceptance', 'Sprint 6 (6A and 6B)', '4'],
            ['Hardening, multi-entity rollout, go-live', 'Sprint 7', '3'],
            ['Hypercare and handover', 'Sprint 8', '2'],
        ],
        col_widths=[7, 4, 2])

    # 2. Assumptions
    heading(doc, '2. Working Assumptions', level=1)
    bullet(doc, 'Sprint length is two weeks by default; extended sprints are explicitly marked.')
    bullet(doc, 'Two engineers are assigned to delivery full-time with ad-hoc data-science input through the sprint.')
    bullet(doc, 'A project manager co-runs ceremonies and stakeholder management at half allocation.')
    bullet(doc, 'The Data Hub endpoint and first historical backfill become available by the end of Sprint 1.')
    bullet(doc, 'Two to three business entities are onboarded at launch.')
    bullet(doc, 'Treasury users are available for user-acceptance sessions during the second half of Sprint 6.')
    bullet(doc, 'Production cloud and database infrastructure are ready by the start of Sprint 7.')

    # 3. Timeline
    heading(doc, '3. Timeline', level=1)

    tl = doc.add_table(rows=3, cols=8)
    tl.style = 'Light Grid Accent 1'
    hdr = tl.rows[0].cells
    labels = ['Sprint 1', 'Sprint 2', 'Sprint 3', 'Sprint 4',
              'Sprint 5', 'Sprint 6', 'Sprint 7', 'Sprint 8']
    weeks = ['Wks 1-3', 'Wks 4-5', 'Wks 6-7', 'Wks 8-9',
             'Wks 10-11', 'Wks 12-15', 'Wks 16-18', 'Wks 19-20']
    themes = ['Integration', 'Preprocessing', 'ML retrain', 'Rules and S7',
              'Reconciliation', 'Recommendation engine', 'Hardening and go-live', 'Hypercare']
    colours = ['D9E8F7', 'D9E8F7', 'CFE4D9', 'CFE4D9',
               'FFF2CC', 'FFF2CC', 'FCE5CD', 'E6E6E6']
    for i in range(8):
        c0 = tl.rows[0].cells[i]; c0.text = ""
        p = c0.paragraphs[0]; r = p.add_run(labels[i]); r.bold = True; r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c0, colours[i])
        c1 = tl.rows[1].cells[i]; c1.text = ""
        p1 = c1.paragraphs[0]; r1 = p1.add_run(weeks[i]); r1.font.size = Pt(9)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c1, colours[i])
        c2 = tl.rows[2].cells[i]; c2.text = ""
        p2 = c2.paragraphs[0]; r2 = p2.add_run(themes[i]); r2.font.size = Pt(9)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c2, colours[i])

    para(doc, '')

    # 4. Milestones
    heading(doc, '4. Delivery Milestones', level=1)
    styled_table(doc,
        ['Milestone', 'Sprint', 'End Week', 'Deliverable'],
        [
            ['Integration proven', 'Sprint 1', 'Week 3', 'Real events flowing through the inbound adapter into the internal event bus.'],
            ['Forecast on real data', 'Sprint 3', 'Week 7', 'Machine-learning models trained on real historical data for each entity.'],
            ['Unified cash position live', 'Sprint 4', 'Week 9', 'Aggregated daily, weekly, and monthly cash view on real data.'],
            ['Cash-accuracy indicator published', 'Sprint 5', 'Week 11', 'First reconciliation summary and composite indicator per entity.'],
            ['Recommendation engine calibrated', 'Sprint 6A', 'Week 13', 'Stress-tested ranked recommendations across all three levers.'],
            ['User-acceptance sign-off', 'Sprint 6B', 'Week 15', 'Treasury users approve recommendation quality and user experience.'],
            ['Production go-live', 'Sprint 7', 'Week 18', 'Platform live in the client cloud, serving real users.'],
            ['Handover complete', 'Sprint 8', 'Week 20', 'Client operations team owns day-to-day operation.'],
        ],
        col_widths=[4.5, 2.3, 1.8, 7.4])

    # 5. Sprint details
    def sprint_section(num, name, weeks, length_note, goal, backlog, acceptance, demo, risks):
        heading(doc, f'5.{num} Sprint {num} — {name}', level=2)
        para(doc, f'Duration: {weeks}  ({length_note})', italic=True)

        para(doc, 'Goal', bold=True)
        para(doc, goal)

        para(doc, 'Backlog', bold=True)
        styled_table(doc,
            ['Item', 'Task'],
            backlog,
            col_widths=[2.5, 13.5])

        para(doc, 'Acceptance Criteria', bold=True)
        for a in acceptance:
            bullet(doc, a)

        para(doc, 'Sprint Demo', bold=True)
        para(doc, demo)

        para(doc, 'Key Risks', bold=True)
        for r in risks:
            bullet(doc, r)

        para(doc, '')

    heading(doc, '5. Sprint-by-Sprint Plan', level=1)

    # ---- Sprint 1 ----
    sprint_section(
        1, 'Data Hub Integration Lock-down', 'Three weeks (Weeks 1 to 3)',
        'Extended by one week to absorb schema-finalisation delays on the integration partner side.',
        'Real envelopes flowing from the Data Hub through the inbound adapter and internal event bus, '
        'with idempotency enforced and a successful historical backfill completed for each entity.',
        [
            ['1.1', 'Finalise the canonical envelope schema with the Data Hub team, including field names, entity identifier semantics, and the envelope identifier source.'],
            ['1.2', 'Exchange the signing key and agree the retry, timeout, and backoff policies for inbound and outbound traffic.'],
            ['1.3', 'Implement adapter changes required to support any agreed schema deltas.'],
            ['1.4', 'Execute a historical backfill covering at least twelve months of events per entity.'],
            ['1.5', 'Validate idempotency against real retries, out-of-order delivery, and partial duplicates.'],
            ['1.6', 'Validate dead-letter-queue behaviour for malformed envelopes, unmapped event types, and handler failures.'],
            ['1.7', 'Produce a data profiling report covering volume, null rates, schema drift, and cardinality per entity.'],
            ['1.8', 'Hold the third week as buffer to absorb schema change requests, access issues, or partner delays.'],
        ],
        [
            'Inbound webhook returns an acknowledgement and writes an event-log row on every valid push.',
            'Duplicate envelopes return a duplicate status without re-processing downstream.',
            'Malformed envelopes land in the dead-letter queue and do not cause non-successful responses.',
            'Historical backfill completes successfully for every entity in scope.',
            'Data profiling report is published and reviewed.',
        ],
        'Live push from the Data Hub produces an event-log row, dispatches to the correct listener, and triggers a partial re-run.',
        [
            'Data Hub schema changes late in the sprint. Mitigation: buffer week plus bulk file fallback.',
            'Historical volume exceeds expectations. Mitigation: batch the backfill and keep streaming the live path separately.',
        ],
    )

    # ---- Sprint 2 ----
    sprint_section(
        2, 'Real-Data Preprocessing and Feature Tables', 'Two weeks (Weeks 4 to 5)',
        'Standard length.',
        'Feature tables computed from real data pass the field-level contract for every entity and '
        'withstand the real-world data quality observed during integration.',
        [
            ['2.1', 'Adjust preprocessing to handle null and missing values on every critical column.'],
            ['2.2', 'Apply currency normalisation across invoices, bills, payments, and contingent inflows.'],
            ['2.3', 'Harmonise timezones and date formats across sources.'],
            ['2.4', 'Map real category values onto the canonical enumerations used by the models.'],
            ['2.5', 'Handle operational edge cases including zero-amount invoices, reversed payments, and credit notes.'],
            ['2.6', 'Re-derive the thin-data threshold per entity based on the observed invoice count distribution.'],
            ['2.7', 'Validate each feature table against the field-level contract defined in the design.'],
            ['2.8', 'Define and implement duplicate and conflicting record merge rules.'],
            ['2.9', 'Produce a per-entity data-quality scorecard covering completeness, validity, and timeliness.'],
        ],
        [
            'All six feature tables build successfully on real data for every entity.',
            'The data-quality scorecard meets the agreed thresholds, or exceptions are documented and accepted.',
            'Thin-data thresholds are committed to configuration per entity.',
            'Field-level contracts pass.',
        ],
        'Side-by-side comparison of real-data feature distributions against the synthetic baseline, per entity.',
        [
            'Data quality worse than anticipated. Mitigation: scorecard-led triage; escalate gaps to the data owners.',
            'Schema misalignment between source systems. Mitigation: raise in the Data Hub integration forum; add transformations inside the adapter.',
        ],
    )

    # ---- Sprint 3 ----
    sprint_section(
        3, 'Machine-Learning Retraining on Real Data', 'Two weeks (Weeks 6 to 7)',
        'Standard length.',
        'All machine-learning modules retrained on real history per entity, with the hierarchical '
        'prior fitted and the model registry seeded to support cold-start prediction and auto-rollback.',
        [
            ['3.1', 'Retrain the receivables prediction models on real history for every entity.'],
            ['3.2', 'Retrain the payables prediction models on real history for every entity.'],
            ['3.3', 'Retrain the credit-risk classifier on real customer history.'],
            ['3.4', 'Investigate and resolve any target leakage surfaced during synthetic-data runs.'],
            ['3.5', 'Fit the hierarchical prior per entity across customer, segment, and global levels.'],
            ['3.6', 'Seed the model registry with real trained artefacts and promote the primary variant.'],
            ['3.7', 'Calibrate the degradation threshold used by the model selector against observed variance.'],
            ['3.8', 'Update the regression-test baselines to match real performance.'],
            ['3.9', 'Verify the primary-to-baseline auto-rollback path end to end.'],
        ],
        [
            'Every machine-learning module is retrained and registered per entity.',
            'The hierarchical prior is fitted, saved, and serves predictions for cold-start customers.',
            'The model registry shows an active primary variant per entity and model.',
            'Automated rollback is verified using a deliberately degraded model.',
            'Baselines are refreshed and committed.',
        ],
        'Live prediction demonstrating the selector routing between primary, baseline, and prior based on customer history depth and recent metrics.',
        [
            'Real data does not support the primary algorithm assumptions. Mitigation: promote the baseline as primary and record the decision in the registry.',
            'Leakage requires feature redesign. Mitigation: minor cases absorbed in Sprint 4; larger cases reported to stakeholders.',
        ],
    )

    # ---- Sprint 4 ----
    sprint_section(
        4, 'Rule-Based Modules and Aggregation Calibration', 'Two weeks (Weeks 8 to 9)',
        'Standard length.',
        'Rule-based modules calibrated on real data, aggregation parameters tuned, and the first '
        'end-to-end task graph run completes successfully on real data for every entity.',
        [
            ['4.1', 'Tune the WIP milestone rules and invoice lag using real project data.'],
            ['4.2', 'Fit the sales pipeline stage probabilities from observed close rates instead of the defaults.'],
            ['4.3', 'Validate the contingent inflows confidence mapping against the real approval-to-payment history.'],
            ['4.4', 'Calibrate the expense forecast lag per category and wire the captured non-PO expenses into the expense input.'],
            ['4.5', 'Tune the aggregation deduplication window and rounding parameters based on observed overlap patterns.'],
            ['4.6', 'Populate source trust baselines from initial metric runs.'],
            ['4.7', 'Execute the first full task graph run on real data for every entity.'],
            ['4.8', 'Sanity-check the resulting cash position against a treasury-owned manual estimate.'],
        ],
        [
            'Every rule-based module produces non-empty, sensible output per entity.',
            'The deduplication engine neither collapses distinct events nor permits true duplicates beyond the agreed tolerance.',
            'Every task graph run completes with a successful status per entity.',
            'Aggregate cash position passes the sanity check against the manual estimate.',
        ],
        'Live unified cash position dashboard showing daily, weekly, and monthly views per entity.',
        [
            'Rule defaults misaligned with reality. Mitigation: addressed within the sprint; all parameters are configuration-driven.',
            'Deduplication too aggressive or too loose. Mitigation: iterative tuning; no code changes needed.',
        ],
    )

    # ---- Sprint 5 ----
    sprint_section(
        5, 'Reconciliation and Cash-Accuracy Indicator', 'Two weeks (Weeks 10 to 11)',
        'Standard length.',
        'Forecast-to-actual reconciliation running per entity, the composite cash-accuracy '
        'indicator published to monitoring, and the regression gate wired into delivery pipelines.',
        [
            ['5.1', 'Wire inbound actual-outcome events from the Data Hub into the reconciliation store.'],
            ['5.2', 'Execute the first monthly reconciliation cycle per entity.'],
            ['5.3', 'Calibrate the days-accuracy target so that variance is meaningful at the observed baseline.'],
            ['5.4', 'Confirm or adjust the weighting between cash accuracy and days accuracy with stakeholders.'],
            ['5.5', 'Emit machine-readable metrics output from every evaluation stage to support the regression gate.'],
            ['5.6', 'Wire the regression gate into delivery pipelines.'],
            ['5.7', 'Build the monitoring dashboard on top of the exposed metrics.'],
            ['5.8', 'Configure alerting on accuracy drift, task-graph failure, dead-letter-queue growth, and ingestion error rate.'],
        ],
        [
            'Realised outcomes are ingested for at least one reconciliation cycle per entity.',
            'The composite indicator is computed per entity and exported to monitoring.',
            'Regression gate runs in the delivery pipeline and either passes or has documented baseline updates.',
            'Monitoring dashboard is reachable and alerts fire against test events.',
        ],
        'Per-entity cash-accuracy indicator alongside a forecast-versus-actual variance visualisation.',
        [
            'Actual-outcome lag longer than expected. Mitigation: extend the reconciliation cadence to six weeks if needed.',
            'Weighting requires re-negotiation. Mitigation: configuration-only change once agreed.',
        ],
    )

    # ---- Sprint 6A ----
    sprint_section(
        '6A', 'Recommendation Engine Calibration on Real Data', 'Two weeks (Weeks 12 to 13)',
        'First half of the four-week recommendation engine sprint.',
        'Recommendation engine recalibrated on real data across all three levers, with scoring weights '
        'stress-tested against a scenario library and cross-module conflicts resolved.',
        [
            ['6A.1', 'Recalibrate the recommendation engine on real cash position and real credit-risk output.'],
            ['6A.2', 'Tune the collections-acceleration lever.'],
            ['6A.3', 'Tune the vendor-deferral lever.'],
            ['6A.4', 'Tune the expense-deferral lever.'],
            ['6A.5', 'Validate policy constraints, including minimum-balance protection and tier-one vendor protection.'],
            ['6A.6', 'Run a scoring-weight sensitivity analysis and document the dominant signals.'],
            ['6A.7', 'Build a scenario library covering quarter-end stress, large invoice slippage, and vendor dispute shocks.'],
            ['6A.8', 'Audit for cross-module conflicts where the engine would advise an action that contradicts a payables or receivables decision.'],
            ['6A.9', 'Hold an internal review to gate progression into user-acceptance sessions.'],
        ],
        [
            'The engine produces non-empty ranked recommendations per entity.',
            'Stress scenarios produce reasonable outputs without violating policy constraints.',
            'Cross-module conflicts are resolved or explicitly accepted with justification.',
            'Internal review approves the progression to user-acceptance testing.',
        ],
        'Walkthrough of the stress scenarios and the sensitivity analysis, demonstrating how scoring weights influence ranking.',
        [
            'Recommendations are too noisy for user review. Mitigation: absorbed through iterative tuning in the sprint.',
            'Cross-module conflicts are harder to resolve than expected. Mitigation: push fixes into the Sprint 7 buffer where necessary.',
        ],
    )

    # ---- Sprint 6B ----
    sprint_section(
        '6B', 'User Acceptance and Feedback Loop', 'Two weeks (Weeks 14 to 15)',
        'Second half of the four-week recommendation engine sprint.',
        'Treasury users review recommendations live, provide feedback, and approve the recommendation '
        'quality and user experience. The feedback loop produces its first scoring-weight proposal.',
        [
            ['6B.1', 'Coordinate four to five user-acceptance sessions with treasury users.'],
            ['6B.2', 'Capture accept, reject, and ignore actions live using the feedback capture endpoint.'],
            ['6B.3', 'Iterate on the captured-expense form based on user feedback.'],
            ['6B.4', 'Apply front-end polish including filtering per entity, Excel export, and daily or weekly toggling.'],
            ['6B.5', 'Run the weight-tuner proposal and review it with stakeholders.'],
            ['6B.6', 'Refine scoring components based on feedback themes.'],
            ['6B.7', 'Document user-acceptance outcomes and obtain sign-off.'],
        ],
        [
            'Every planned user-acceptance session is delivered.',
            'At least eighty percent of feedback items are addressed or explicitly backlogged with rationale.',
            'Treasury users sign off on recommendation quality and user experience.',
            'A first weight-tuner proposal is generated and reviewed.',
        ],
        'Treasury user walks through recommendations live, provides feedback, and reviews the resulting scoring-weight proposal.',
        [
            'User availability misaligned with the sprint window. Mitigation: book sessions during Sprint 4; keep recorded walkthrough as a fallback.',
            'Feedback volume insufficient for meaningful tuning. Mitigation: expected for the first pass; the tuner remains advisory.',
        ],
    )

    # ---- Sprint 7 ----
    sprint_section(
        7, 'Hardening, Multi-Entity Rollout, and Go-Live', 'Three weeks (Weeks 16 to 18)',
        'Extended by one week to absorb security and performance findings ahead of cut-over.',
        'Platform live in the client cloud with real users, all entities configured, security and '
        'performance validated, alerts and runbooks in place.',
        [
            ['7.1', 'Configure every entity in the production environment.'],
            ['7.2', 'Apply per-entity configuration overrides where needed.'],
            ['7.3', 'Implement an automated entity-isolation test in the delivery pipeline.'],
            ['7.4', 'Execute the security review, including a penetration test on the inbound webhook and a role-based access audit.'],
            ['7.5', 'Execute performance testing covering the application programming interface and concurrent task-graph runs.'],
            ['7.6', 'Wire alerting on accuracy drift, task-graph failure, dead-letter-queue growth, and ingestion error rate.'],
            ['7.7', 'Publish runbooks for dead-letter replay, feature-version freezing, model rollback, and emergency stop.'],
            ['7.8', 'Reserve the third week as buffer to absorb security fixes, performance tuning, and last-minute configuration changes.'],
            ['7.9', 'Execute the go-live cut-over plan.'],
        ],
        [
            'All entities are configured and isolation-tested in production.',
            'Security findings classified as critical or high are resolved prior to go-live.',
            'Performance testing meets the agreed service-level targets.',
            'All runbooks are reviewed by the client operations team.',
            'Cut-over is signed off and a rollback procedure is in place.',
        ],
        'Live production environment with real users exercising the application programming interface and dashboard, including live monitoring and alerting.',
        [
            'Security review surfaces critical findings. Mitigation: buffer week absorbs fixes; contingency is a one-week delay of cut-over.',
            'Performance under real load below expectations. Mitigation: tune database pooling, add caching where necessary.',
        ],
    )

    # ---- Sprint 8 ----
    sprint_section(
        8, 'Hypercare and Handover', 'Two weeks (Weeks 19 to 20)',
        'Standard length. No new features.',
        'Stabilise the production environment, handle day-one issues, hand over operational '
        'ownership to the client operations team, and deliver the close-out report.',
        [
            ['8.1', 'Hold daily stand-ups with the client operations team.'],
            ['8.2', 'Triage and fix bugs surfaced by real usage.'],
            ['8.3', 'Tune the cash-accuracy indicator based on the first weeks of real operation.'],
            ['8.4', 'Review and, where applicable, act on the second weight-tuner proposal.'],
            ['8.5', 'Validate the first volume-triggered retraining end to end.'],
            ['8.6', 'Deliver handover sessions covering operations, incident response, and routine tuning.'],
            ['8.7', 'Publish the close-out report including baselines, lessons learned, and a forward roadmap.'],
            ['8.8', 'Hold the final retrospective with the client.'],
        ],
        [
            'No critical or high bugs remain open at the end of the sprint.',
            'The client operations team can independently execute a full task-graph run, a reconciliation, and a dead-letter replay.',
            'The close-out report is delivered and accepted.',
        ],
        'Final retrospective with the client summarising the first weeks of real operation. No standalone demo needed as the production system has been live for two weeks.',
        [
            'Critical issue surfaces in the final week. Mitigation: budget one extra week of hypercare if required.',
            'Operations team is not yet confident. Mitigation: schedule additional handover sessions during the first week of hypercare.',
        ],
    )

    # 6. Risk register
    heading(doc, '6. Risk Register', level=1)
    styled_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['Data Hub delivery slips', 'Medium', 'High',
             'Buffer week in Sprint 1; bulk file backfill in parallel until the live webhook stabilises.'],
            ['Real data quality worse than the design assumes', 'Medium', 'High',
             'Sprint 2 is dedicated to this; a per-entity data-quality scorecard surfaces gaps early.'],
            ['User-acceptance users unavailable on schedule', 'Medium', 'Medium',
             'Confirm availability in Sprint 4; recorded walkthrough as a fallback.'],
            ['Hierarchical prior underperforms on real data', 'Low to Medium', 'Medium',
             'Graceful degradation is already built in across prior, baseline, and primary variants.'],
            ['Multi-entity configuration drift between test and production', 'Low', 'Medium',
             'Automated entity-isolation test in the delivery pipeline before go-live.'],
            ['Security review surfaces critical findings', 'Medium', 'High',
             'Buffer week in Sprint 7; one-week cut-over delay held in reserve.'],
            ['Recommendations rejected by users during acceptance', 'Medium', 'High',
             'Sprint 6A scenario library and sensitivity analysis catches issues before the acceptance block.'],
            ['Client operations not ready for handover', 'Low', 'Medium',
             'Handover sessions begin on day one of Sprint 8; additional sessions reserved.'],
            ['Indicator weighting disagreement between engineering and business', 'Low', 'Low',
             'Configuration-only change after the first reconciliation reading.'],
            ['Production performance below service-level target', 'Medium', 'Medium',
             'Performance test in Sprint 7; database pooling and caching held as levers.'],
        ],
        col_widths=[5.5, 2, 2, 6.5])

    # 7. Dependencies
    heading(doc, '7. External Dependencies', level=1)
    styled_table(doc,
        ['Dependency', 'Needed By', 'Impact If Delayed'],
        [
            ['Data Hub endpoint, signing key, and agreed envelope schema',
             'End of Sprint 1',
             'Downstream sprints slip by the delay duration; synthetic bulk file fallback reduces, but does not eliminate, the impact.'],
            ['Historical backfill covering twelve months per entity',
             'End of Sprint 1',
             'Machine-learning retraining in Sprint 3 cannot begin; rule calibration in Sprint 4 operates on thinner data.'],
            ['Approved treasury-owned configuration tables (vendor priority, escalation rules)',
             'End of Sprint 4',
             'Rule-based modules fall back to defaults; recommendation engine calibration in Sprint 6 is less faithful.'],
            ['Availability of treasury users for user-acceptance sessions',
             'Sprint 6B',
             'User-acceptance sign-off delays; go-live must be pushed right.'],
            ['Production cloud and database infrastructure',
             'Start of Sprint 7',
             'Hardening and go-live cannot begin.'],
            ['Security review partner',
             'Mid-Sprint 7',
             'Cut-over slips to the buffer week, reducing hypercare runway.'],
        ],
        col_widths=[6, 3, 7])

    # 8. Governance
    heading(doc, '8. Governance and Communication', level=1)
    styled_table(doc,
        ['Forum', 'Frequency', 'Participants', 'Purpose'],
        [
            ['Daily engineering stand-up', 'Daily', 'Engineering team', 'Blockers, progress, immediate decisions.'],
            ['Sprint review and demo', 'End of every sprint', 'Engineering team and client stakeholders', 'Demonstrate the sprint deliverable and gain acceptance.'],
            ['Sprint retrospective', 'End of every sprint', 'Engineering team', 'Process improvement for the next sprint.'],
            ['Client steering committee', 'Every two weeks', 'Engineering lead, project manager, client lead', 'Risks, scope, schedule, resource escalations.'],
            ['Hypercare stand-up', 'Daily during Sprint 8', 'Engineering team and client operations team', 'Triage incidents, track handover.'],
        ],
        col_widths=[4.5, 2.8, 4.5, 4.2])

    # 9. Definition of Done
    heading(doc, '9. Sprint Definition of Done', level=1)
    para(doc, 'A sprint is complete when all of the following hold:')
    bullet(doc, 'Every acceptance criterion in the sprint is met.')
    bullet(doc, 'Every committed backlog item is merged and tested, or explicitly carried over with a stated reason.')
    bullet(doc, 'The sprint review and demo have been delivered.')
    bullet(doc, 'No critical or high defects remain open.')
    bullet(doc, 'Documentation relevant to the sprint deliverable has been updated.')

    # 10. Scope controls
    heading(doc, '10. Scope Controls', level=1)
    para(doc,
         'The plan is fixed-scope with in-sprint flexibility on task ordering. Scope changes '
         'beyond the items listed are absorbed by the next sprint only if they do not endanger '
         'the go-live date. Larger scope changes require explicit re-planning at the steering '
         'committee and may move the go-live date. The plan explicitly holds buffer in Sprint 1, '
         'Sprint 6, and Sprint 7 to absorb uncertainty rather than to accommodate new scope.')

    # 11. Summary
    heading(doc, '11. Summary', level=1)
    para(doc,
         'Twenty weeks. Eight sprints. Four weeks of explicit buffer allocated to the highest-risk '
         'phases. One production cut-over at the end of Sprint 7. Two weeks of hypercare to hand '
         'ownership to the client operations team.')

    out = "Project_Plan.docx"
    doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    build_system_architecture()
    build_project_plan()
